from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE = Path(r"D:/Диплом что-то адекватное/charts-for-diplom/текст")
SRC_MD = BASE / "Отчет_НИР_Байрамова_по_образцу.md"
OUT_DOCX = BASE / "Отчет_НИР_Байрамова_по_образцу.docx"
IMG1 = BASE / "assets/custom/scheme_app_architecture.png"
IMG2 = BASE / "assets/custom/scheme_data_flow.png"


def add_heading_like(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(text)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(14)
    else:
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(13)
    doc.add_paragraph("")


def add_plain_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_image_with_caption(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Pt(430))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    doc.add_paragraph("")


def build_docx() -> None:
    text = SRC_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    in_description = False
    inserted_images = False

    for raw in lines:
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line == "---":
            doc.add_page_break()
            continue

        if line.startswith("# "):
            add_heading_like(doc, line[2:].strip(), level=1)
            continue

        if line.startswith("## "):
            add_heading_like(doc, line[3:].strip(), level=2)
            in_description = line[3:].strip().upper() == "ОПИСАНИЕ РАБОТЫ"
            continue

        if line.startswith("- "):
            add_plain_paragraph(doc, f"• {line[2:].strip()}")
            continue

        if line[:2].isdigit() and line[1:3] == ". ":
            add_plain_paragraph(doc, line)
            continue

        add_plain_paragraph(doc, line)

        if in_description and "В результате выполнения работы создан прототип" in line and not inserted_images:
            add_image_with_caption(doc, IMG1, "Рис. 1 Архитектура разработанного приложения")
            add_image_with_caption(doc, IMG2, "Рис. 2 Поток данных при построении профилей и трендов")
            inserted_images = True

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
